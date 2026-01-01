import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from deep_translator import GoogleTranslator
import os
import time
from PIL import Image
import pytesseract
import platform

# Set Tesseract path based on operating system
def setup_tesseract():
    """Configure Tesseract OCR path based on the operating system"""
    system = platform.system()
    if system == "Windows":
        # Common Windows installation paths
        possible_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in possible_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                return path
        # If not found, try to use system PATH
        print("警告: 未找到 Tesseract OCR，请确保已安装并添加到系统 PATH")
    # For Linux/Mac, assume tesseract is in PATH
    return None

setup_tesseract()

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF, handling both text and images with OCR"""
    doc = fitz.open(pdf_path)
    all_text = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Extract text directly
        text = page.get_text()
        
        if text.strip():
            all_text.append({
                'page': page_num + 1,
                'text': text.strip(),
                'type': 'text'
            })
        else:
            # If no text found, try OCR on the page
            print(f"  检测到扫描图片页面 {page_num + 1}，正在进行OCR识别...")
            
            # Convert page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution for better OCR
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            try:
                # Perform OCR with English and Chinese language support
                ocr_text = pytesseract.image_to_string(img, lang='eng+chi_sim')
                
                if ocr_text.strip():
                    all_text.append({
                        'page': page_num + 1,
                        'text': ocr_text.strip(),
                        'type': 'ocr'
                    })
                    print(f"  ✓ OCR识别成功，提取了 {len(ocr_text)} 个字符")
                else:
                    all_text.append({
                        'page': page_num + 1,
                        'text': '[扫描图片页面 - OCR未能识别到文本]',
                        'type': 'image'
                    })
            except Exception as e:
                print(f"  ✗ OCR识别失败: {str(e)}")
                all_text.append({
                    'page': page_num + 1,
                    'text': f'[扫描图片页面 - OCR错误: {str(e)}]',
                    'type': 'image'
                })
    
    doc.close()
    return all_text

def translate_text(text, max_retries=3):
    """Translate text from English to Chinese with retry logic"""
    if not text or text.strip() == '' or '[扫描图片页面' in text:
        return text
    
    for attempt in range(max_retries):
        try:
            # Translate to Chinese using deep-translator
            translator = GoogleTranslator(source='auto', target='zh-CN')
            translation = translator.translate(text)
            time.sleep(0.5)  # Add delay to avoid rate limiting
            return translation
        
        except Exception as e:
            print(f"Translation attempt {attempt + 1} failed: {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(2)  # Wait before retrying
            else:
                return f"[翻译失败: {text}]"
    
    return text

def create_bilingual_docx(pdf_path, output_path):
    """Create a bilingual DOCX file from PDF"""
    print(f"\n处理文件: {pdf_path}")
    
    # Extract text from PDF
    print("正在提取PDF内容...")
    contents = extract_text_from_pdf(pdf_path)
    
    if not contents:
        print(f"警告: {pdf_path} 中没有找到文本内容")
        return False
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(os.path.basename(pdf_path).replace('.pdf', ''), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process each page
    for idx, content in enumerate(contents):
        print(f"处理第 {content['page']} 页...")
        
        # Add page number
        page_heading = doc.add_heading(f"第 {content['page']} 页 / Page {content['page']}", level=2)
        
        original_text = content['text']
        
        if content['type'] == 'image':
            # For scanned images that couldn't be processed, add a note
            para = doc.add_paragraph()
            run = para.add_run(original_text)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.italic = True
        elif content['type'] == 'ocr':
            # For OCR-extracted text, add a note indicating it was OCR processed
            para = doc.add_paragraph()
            run = para.add_run("【注意：此页面通过OCR识别提取】\n")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 140, 0)  # Orange color
            run.italic = True
            
            # Split into paragraphs
            paragraphs = original_text.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Add original text
                para = doc.add_paragraph()
                run = para.add_run("【英文原文】\n")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 139)
                para.add_run(para_text.strip())
                
                # Translate and add Chinese text
                print(f"  正在翻译OCR识别的段落...")
                chinese_text = translate_text(para_text.strip())
                
                para = doc.add_paragraph()
                run = para.add_run("【中文翻译】\n")
                run.bold = True
                run.font.color.rgb = RGBColor(139, 0, 0)
                para.add_run(chinese_text)
                
                # Add separator
                doc.add_paragraph("─" * 50)
        else:
            # Split into paragraphs
            paragraphs = original_text.split('\n\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                
                # Add original English text
                para = doc.add_paragraph()
                run = para.add_run("【英文原文】\n")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 139)
                para.add_run(para_text.strip())
                
                # Translate and add Chinese text
                print(f"  正在翻译段落 {idx + 1}...")
                chinese_text = translate_text(para_text.strip())
                
                para = doc.add_paragraph()
                run = para.add_run("【中文翻译】\n")
                run.bold = True
                run.font.color.rgb = RGBColor(139, 0, 0)
                para.add_run(chinese_text)
                
                # Add separator
                doc.add_paragraph("─" * 50)
        
        # Add page break after each page (except the last one)
        if idx < len(contents) - 1:
            doc.add_page_break()
    
    # Save document
    print(f"正在保存文件: {output_path}")
    doc.save(output_path)
    print(f"✓ 完成: {output_path}")
    return True

def main():
    """Main function to process all PDF files in current directory"""
    current_dir = os.getcwd()
    pdf_files = [f for f in os.listdir(current_dir) if f.lower().endswith('.pdf')]
    
    if not pdf_files:
        print("未找到PDF文件")
        return
    
    print(f"找到 {len(pdf_files)} 个PDF文件")
    print("=" * 60)
    
    success_count = 0
    failed_files = []
    
    for pdf_file in pdf_files:
        pdf_path = os.path.join(current_dir, pdf_file)
        docx_file = pdf_file.replace('.pdf', '.docx')
        docx_path = os.path.join(current_dir, docx_file)
        
        try:
            if create_bilingual_docx(pdf_path, docx_path):
                success_count += 1
            else:
                failed_files.append(pdf_file)
        except Exception as e:
            print(f"✗ 处理 {pdf_file} 时出错: {str(e)}")
            failed_files.append(pdf_file)
    
    print("\n" + "=" * 60)
    print(f"处理完成!")
    print(f"成功: {success_count}/{len(pdf_files)}")
    
    if failed_files:
        print(f"失败的文件: {', '.join(failed_files)}")
    
    print("\n注意: 已启用OCR功能，可以处理扫描图片的PDF文件")
    print("Tesseract OCR 安装路径: C:\\Program Files\\Tesseract-OCR\\tesseract.exe")

if __name__ == "__main__":
    main()
